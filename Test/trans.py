# -*- coding:UTF-8 -*-
from random import choice
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import load_workbook

def xlsx2docx(fn):
    #打开excel文件，如有公式，读取公式计算结果
    wb=load_workbook(fn,data_only=True)
    #创建空白word 文件
    document=Document()

    #查看所有可用的表格样式
    table_styles=[style for style in document.styles 
                if style.type==WD_STYLE_TYPE.TABLE]
    #print(talbe_styles)

    #遍历EXCEL 文件中所有的worksheet
    for ws in wb.worksheets:
        rows=list(ws.rows)
        #增加段落，也就是表格的名称
        document.add_paragraph(ws.title)
        #根据worKsheet的行数和列数，在word 文件中创建核实大小的表格

        mylen = 0

        for i in rows:
            # print(len(i))
            mylen = len(i)

        table=document.add_table(rows=len(rows),
                                cols= mylen,
                                style=choice(table_styles))
        #从worksheet读取数据，写入word文件中的表格
        for  irow,row in enumerate(rows):
            for icol,col in enumerate(row):
                if col.value == None:
                    continue
                s = str( col.value )

                if 'X' in s:
                    print(s)
                # print(str(col.value))
                # table.cell(irow,icol).text=str(col.value)
    #保存word文件
    # document.save(fn[:-4]+'docx')
#调用函数，进行数据导入
xlsx2docx(r'C:\Users\HP\Desktop\abc\顺北5井基本数据表.xlsx')