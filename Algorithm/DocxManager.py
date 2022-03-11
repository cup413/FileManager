import os
import os.path as pt

from win32com import client as wc
import time

import docx

import pandas as pd

import threading

import os

import re



class DocxManager:

    def toolDropSpace(self, s):
        tmps = ''
        for i in s:
            if i!= ' ':
                tmps = tmps+i
        return tmps

    def makeDoc(self, path):
        '''
        获取path所指定的word文档的文字和表格，分别保存在
        self.text和self.tables里面
        :param path:
        :return:
        '''
        doc = docx.Document(path)

        pa = [p.text for p in doc.paragraphs]
        self.text = ''.join(pa)

        self.tables = doc.tables


    def checkInTables(self):
        '''
        检查word文件是否含有地层分层表，
        判断方法是cell(0,0)含有地层
        cell(1,0)含有界或者系
        :return:
        '''
        for table in self.tables:
            try:
                s1 = table.cell(0,0).text
                s2 = table.cell(1,0).text
                s3 = table.cell(0,1).text
            except:
                continue

            if '地' in s1 and ( '界' in s2 or '系' in s2):
                return True

        return False

    def returnTable(self):
        '''
        获取基本信息表
        :return:
        '''
        mytable = []
        for table in self.tables:
            col = len(table.columns)

            #             print( table.cell(0,0).text )
            if ('位置' not in table.cell(0, 0).text) and table.cell(0, 0).text != '基本数据':
                continue

            for i in range(len(table.rows)):

                atable = []
                for j in range(col):
                    atable.append(table.cell(i, j).text)
                mytable.append(atable)
        # print('=============')
        return mytable

    def getInfoFromTable(self):
        '''
        从基本信息表获取数据，
        包括大地坐标、经纬度、地面海拔、补心海拔和补心高
        :return:
        '''
        def addData(p, i, j, atable):
            # print(p,i,j, atable[i][j])
            while (atable[i][j] == atable[i][j + 1]):
                j = j + 1
            if self.info[p] == 'T':
                self.info[p] = atable[i][j + 1]

        table = self.returnTable()

        self.info = ['T' for i in range(7)]
        for i in range(len(table)):
            for j in range(len(table[i])):
                if ('X' in table[i][j]) and ('纵' in table[i][j]):
                    addData(0, i, j, table)
                #                     while( table[i][j] ==  table[i][j+1] ):
                #                         j = j + 1
                #                     if self.info[0] == 'T':
                #                         self.info[0] = table[i][j+1]
                if ('Y' in table[i][j]) and ('横' in table[i][j]):
                    addData(1, i, j, table)
                #                     while( table[i][j] ==  table[i][j+1] ):
                #                         j = j + 1
                #                     if self.info[1] == 'T':
                #                         self.info[1] = table[i][j+1]
                if table[i][j] == '东经':
                    addData(2, i, j, table)
                #                     if self.info[2] == 'T':
                #                         self.info[2] = table[i][j+1]
                if table[i][j] == '北纬':
                    addData(3, i, j, table)
                #                     if self.info[3] == 'T':
                #                         self.info[3] = table[i][j+1]
                if table[i][j] == '地面海拔':
                    addData(4, i, j, table)
                #                     self.info[4] = table[i][j+1]
                if table[i][j] == '补心海拔' or table[i][j] == '补芯海拔':
                    addData(5, i, j, table)
                #                     self.info[5] = table[i][j+1]
                if table[i][j] == '补心高':
                    addData(6, i, j, table)
        #                     self.info[6] = table[i][j+1]
        lst = ['X', 'Y', 'Lat', 'Lon', 'Hd', 'Hb', 'Hbb']
        for i in range(len(lst)):
            print(lst[i], self.info[i])

        if self.info[0] != 'T':
            return True
        return False

    def getInfoFromText(self):
        '''
        从文字中获取基本信息
        包括大地坐标、经纬度、地面海拔、补心海拔和补心高
        :return:
        '''
        def delSpace(s):
            tmps = ''
            for i in s:
                if i != ' ':
                    tmps = tmps + i
            return tmps

        def getNum(pat, i):
            reg = pat
            s = re.search(reg, self.text)
            print(s)
            s = s.group(0)
            #             print(s)
            s = delSpace(s)
            reg = '[0-9]*\.[0-9]*'
            s = re.search(reg, s)
            self.info[i] = s.group(0)

        self.info = ['T' for i in range(7)]

        getNum(r'纵.{0,4}X.*?m', 0)
        getNum(r'横.{0,4}Y.*?m', 1)
        #         getNum(r'东')
        getNum(r'地面海拔.*?m', 4)
        getNum(r'补.{1}海拔.*?m', 5)
        getNum(r'补心高.*?m', 6)

        lst = ['X', 'Y', 'Lat', 'Lon', 'Hd', 'Hb', 'Hbb']
        for i in range(len(lst)):
            print(lst[i], self.info[i])

        if self.info[0] !='T':
            return True
        else:
            return False

    def returnLst(self):
        '''
        根据地层-界（系）找到地层分层数据
        并且靶这个分层数据返回（lst)
        :return:
        : lst, 一个二维数组，保存分层数据表
        '''
        # print('==========find layer================')
        flst = []
        for table in self.tables:
            try:
                # print('aaaaaaaaaaaaaaabbbbbbb')
                lst = []
                #     table = tables[1]#获取文件中的第9个表格
                col = len(table.columns)

                # print(table.cell(0, 0).text)
                s = ''
                for i in table.cell(0, 0).text:
                    if i != ' ':
                        s = s + i
                try:
                    table.cell(1,0).text
                except:
                    continue


                # print('333333333333333333333333')
                s2 = table.cell(1,0).text
                if ('地层' in s) and ( '界' in s2 or '系' in s2 ):

                    for i in range(1, len(table.rows)):  # 从表格第二行开始循环读取表格数据

                        alst = []
                        for j in range(col):
                            alst.append(table.cell(i, j).text)
                            # print(table.cell(i, j).text, end="  ")
                        lst.append(alst)
                        # print(alst)
                    # print('===========lst over')
                    # print('======lst')
                    # print(lst)
                    # print(len(lst), len(flst))
                    if len(lst) > len(flst):
                        # print('-----------------')
                        flst = lst
            except:
                pass

        # print('aaaaaa   flst')
        # print(flst)
        return flst

    def getLayer(self):
        '''
        从获得的分层数据数组中提取系-底界深度
        :return:
        : ret，一个二维数组，保存系-底界深度
        '''
        def getDeepth(s):
            if ('底' in s and '深' in s) or ('井' in s and '深' in s):
                return True
            return False

        def getXi(s):
            if '系' in s:
                return True
            return False

        def getidx(lst, isTrue):
            # print(lst[0])
            idx = 0
            for i in range(len(lst[0])):

                if isTrue(lst[0][i]):
                    idx = i
                    break
            return idx
        lst = self.returnLst()

        # print('========lst')
        # print(lst)

        # idx = 0
        # for i in range(len(lst[0])):
        #
        #     if ('底' in lst[0][i] and '深' in lst[0][i]) or '井' in lst[0][i] and '深' in lst[0][i]:
        #         idx = i
        #         break
        # idxXi = 0
        idx = getidx(lst, getDeepth)
        idxXi = getidx(lst, getXi)
        # print('idx, idxXi', idx, idxXi)

        ret = []
        for i in lst:
            #             print(i[0],i[1], i[idx])
            ret.append([i[idxXi], i[idx]])

        return ret

    def getLayer1(self):

        '''
        将getLayer()获得的系-底界深度去重得到最后的结果
        :return:
        : self.Layer，二维数组，保存系-底界深度
        '''
        if not self.checkInfo() and not self.checkInTables():
            print('地层不存在')
            return False

        ret = self.getLayer()

        myret = []
        for i in ret:
            if i[1] == '':
                pass
            else:
                myret.append(i)
        ret = myret

        alreadyIn = set()
        fret = []
        for i in range(len(ret) - 1, -1, -1):
            if ret[i][0] in alreadyIn:
                continue

            alreadyIn.add(ret[i][0])
            fret.append([ret[i][0], ret[i][1]])
            print([ret[i][0], ret[i][1]])
        # print('==============')
        fret.reverse()

        self.layer = fret

        if len(self.layer) == 0:
            return False
        return True

    def saveLayer(self, name):
        '''
        保存地层分层数据
        :param name: 保存分层数据的名字
        :return:
        '''
        if self.layer == '':
            return False

        folder = r'D:\李晨星文件夹\项目文件\塔里木程小桂\layer\%s.csv'
        pth = folder % name


        f = pd.DataFrame(self.layer[1:], columns=['层', '底界深度'])
        f.to_csv(pth, index=False)
        print(pth, '   保存成功')
        return True

    def saveInfo(self, name):
        '''
        保存基本信息
        :param name: 基本信息的名字
        :return:
        '''
        if self.info[0] == 'T':
            print('info 为空，不能保存')
            return False

        folder = r'D:\李晨星文件夹\项目文件\塔里木程小桂\info\%s.csv'
        pth = folder % name
        print(pth,'   保存成功')

        col = ['X', 'Y', 'Lat', 'Lon', 'Hd', 'Hb', 'Hbb']
        f = pd.DataFrame([self.info], columns=col)
        f.to_csv(pth, index=False)

        return True


    def checkInfo(self):
        '''
        从self.text检查word文档是否包含基本信息和分层数据
        :return:
        '''
        print('===============检查信息===========')
        if '座标' in self.text or '坐标' in self.text:
            print('坐标存在')

        if '补心' in self.text or '补芯' in self.text:
            print('补心存在')

        if '地层分层' in self.text:
            print('地层分层存在')

            return True

        if '基本数据' in self.text:
            print('基本数据存在')

        if '东经' in self.text:
            print('经纬度存在')

        return False
